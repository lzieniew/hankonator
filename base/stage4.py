from docx.shared import Pt


class Stage4:

    def __init__(self, erd):
        self.erd = erd

        self.rules = {}

        self.populate_rules()

    def populate_rules(self):
        for entity in self.erd.entities:
            self.rules[entity.name_singular] = []
            relationships = self.erd.get_relationships_connected_wit_entity(entity.name_singular)
            for rel in relationships:
                if rel.get_other_ends_multiplicity(entity.name_singular)[0] == '0':
                    self.rules[entity.name_singular].append(entity.name_singular + ' nie musi być powiązany z żadnym ' + rel.get_other_entity_name(entity.name_singular) + '\n')
                elif rel.get_other_ends_multiplicity(entity.name_singular)[0] == '1':
                    self.rules[entity.name_singular].append(entity.name_singular + ' musi być powiązany z przynajmniej jednym ' + rel.get_other_entity_name(entity.name_singular) + '\n')

                if rel.get_other_ends_multiplicity(entity.name_singular)[-1] == '1':
                    self.rules[entity.name_singular].append(entity.name_singular + ' jest powiązany z maksymalnie jednym ' + rel.get_other_entity_name(entity.name_singular) + '\n')
                elif rel.get_other_ends_multiplicity(entity.name_singular)[-1] == 'N':
                    self.rules[entity.name_singular].append(entity.name_singular + ' jest powiązany z wieloma ' + rel.get_other_entity_name(entity.name_singular) + '\n')
        print(self.rules)


    def build(self, document):

        header = document.add_paragraph()
        header.add_run('4. Reguły funkcjonowania').font.size = Pt(24)
        header.add_run().add_break()

        counter = 1
        rules_counter = 1

        for entity in self.rules.keys():
            rules_paragraph = document.add_paragraph()
            rules_paragraph.keep_together = True

            if self.rules[entity]:
                rules_paragraph.add_run(str(counter) + '. Reguły dla KAT/' + '{0:03}'.format(self.erd.get_entity_by_name(entity).id) + ' ' + entity).font.size = Pt(16)
            rules_paragraph.add_run().add_break()
            for rule in self.rules[entity]:
                rules_paragraph.add_run('REG/' + '{0:03}'.format(rules_counter)).font.bold=True
                rules_paragraph.add_run('\t\t')
                rules_paragraph.add_run(rule)

                rules_counter += 1

            counter += 1
        document.add_page_break()

        print('Etap 4 wygenerowany')
