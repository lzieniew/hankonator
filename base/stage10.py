from docx.shared import Pt

from generation import Project

class Stage10(object):

    def __init__(self, erd):
        self.erd = erd

    def add_foreign_keys(self):
        for entity in self.erd.entities:
            relationships = self.erd.get_relationships_connected_wit_entity(entity.name_singular)
            for relationship in relationships:
                if relationship.get_this_ends_multiplicity(entity.name_singular)[-1] == 'N':
                    other_entity = self.erd.get_entity_by_name(relationship.get_other_entity_name(entity.name_singular))
                    entity.foreign_keys.append(other_entity.get_key())



    def build(self, document):
        header = document.add_paragraph()
        header.add_run('10. Transformacja modelu konceptualnego do modelu logicznego').font.size = Pt(Project.HEADER_SIZE)
        header.add_run().add_break()

        for relationship in self.erd.relationships:
            relationship_paragraph = document.add_paragraph()
            relationship_paragraph.paragraph_format.keep_together = True
            relationship_paragraph.add_run('ZWI' + '{0:03}'.format(relationship.id)).bold = True
            relationship_paragraph.add_run(' ' + relationship.name + '(' + relationship.left_entity.upper() + '(' + relationship.left_quantity + '):' + relationship.right_entity.upper() + '(' + relationship.right_quantity + ')')
            relationship_paragraph.add_run().add_break()

            left_entity = self.erd.get_entity_by_name(relationship.left_entity)
            right_entity = self.erd.get_entity_by_name(relationship.right_entity)
            left_entity.foreign_keys = []
            right_entity.foreign_keys = []
            relationship_paragraph.add_run('ENC/' + '{0:03}'.format(left_entity.id) + ' ' + left_entity.name_singular.upper() + ' ')
            left_entity.build_argument_list(relationship_paragraph)
            relationship_paragraph.add_run().add_break()
            relationship_paragraph.add_run('ENC/' + '{0:03}'.format(right_entity.id) + ' ' + right_entity.name_singular.upper() + ' ')
            right_entity.build_argument_list(relationship_paragraph)
            relationship_paragraph.add_run().add_break()
            relationship_paragraph.add_run().add_break()
            relationship_paragraph.add_run('Po przekształceniu otrzymujemy')
            relationship_paragraph.add_run().add_break()
            relationship_paragraph.add_run().add_break()

            self.add_foreign_keys()

            relationship_paragraph.add_run(
                'REL/' + '{0:03}'.format(left_entity.id) + ' ' + left_entity.name_plural + ' ')
            left_entity.build_argument_list(relationship_paragraph)
            relationship_paragraph.add_run().add_break()
            relationship_paragraph.add_run(
                'REL/' + '{0:03}'.format(right_entity.id) + ' ' + right_entity.name_plural + ' ')
            right_entity.build_argument_list(relationship_paragraph)


        print('Etap 10 wygenerowany!')