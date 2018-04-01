from docx.shared import Pt

from generation import Project

class Stage8(object):

    def __init__(self, erd):
        self.erd = erd

    def build(self, document):
        header = document.add_paragraph()
        header.add_run('8. Definicje predykatowe encji i związków').font.size = Pt(Project.HEADER_SIZE)
        header.add_run().add_break()

        entities_paragraph = document.add_paragraph()
        entities_paragraph.add_run('8.1. Encje').font.size = Pt(Project.SECONDAR_HEADER_SIZE)
        entities_paragraph.add_run().add_break()

        for entity in self.erd.entities:
            entities_paragraph.add_run('ENC/' + '{0:03}'.format(entity.id) + ' ' + entity.name_singular.upper() + ' ').bold = True
            entities_paragraph.add_run(entity.repr_attributes()).italic = True
            entities_paragraph.add_run().add_break()

        document.add_page_break()

        print(u'Etap 8 wygenerowany!')
