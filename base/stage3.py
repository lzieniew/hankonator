from docx.shared import Pt


class Stage3:

    def __init__(self, erd):
        self.erd = erd

    def build(self, document):

        header = document.add_paragraph()
        header.add_run('3. Kategorie').font.size = Pt(24)
        header.add_run().add_break()

        for entity in self.erd.entities:
            entity_paragraph = document.add_paragraph()
            entity_paragraph.keep_together = True
            entity_paragraph.add_run('KAT/' + '{0:03}'.format(entity.id) + ' ' + entity.name_singular).font.size = Pt(16)
            entity_paragraph.add_run().add_break()
            entity_paragraph.add_run('\tOpis: ').bold = True
            entity_paragraph.add_run('[Tutaj krótki opis encji]')
            entity_paragraph.add_run().add_break()
            entity_paragraph.add_run('\tAtrybuty:').bold = True
            entity_paragraph.add_run().add_break()
            for attribute in entity.attributes:
                entity_paragraph.add_run('\t\t' + repr(attribute.name) + ' - [Opis atrybutu]')
                entity_paragraph.add_run().add_break()

        print('Etap 3 wygenerowany')
