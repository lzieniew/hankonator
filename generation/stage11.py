from random import randint, randrange

from docx.shared import Pt

from generation import Project


class Stage11(object):

    TABLE_FONT_SIZE = 6
    EXAMPLES_COUNT = 4

    def __init__(self, erd):
        self.erd = erd

    def build(self, document):
        header = document.add_paragraph()
        header.add_run('11. Definicje schematów relacji i przykładowe dane w poszczególnych tabelach').font.size = Pt(Project.HEADER_SIZE)
        header.add_run().add_break()

        for entity in self.erd.entities:
            relation_paragraph = document.add_paragraph()
            relation_paragraph.add_run('REL/' + '{0:03}'.format(entity.id) + ' ' + entity.name_plural + '/' + entity.name_singular.upper()).bold = True
            relation_paragraph.add_run().add_break()
            relation_paragraph.add_run('Opis schematu relacji:')
            relation_paragraph.add_run().add_break()

            table = document.add_table(rows = len(entity.attributes) + len(entity.foreign_keys) + 1, cols = 10)
            hdr_cells = table.rows[0].cells
            runs = []
            runs.append(hdr_cells[0].paragraphs[0].add_run('Nazwa atrybutu'))
            runs.append(hdr_cells[1].paragraphs[0].add_run('Dziedzina'))
            runs.append(hdr_cells[2].paragraphs[0].add_run('Maska'))
            runs.append(hdr_cells[3].paragraphs[0].add_run('OBL'))
            runs.append(hdr_cells[4].paragraphs[0].add_run('Wart. Dom.'))
            runs.append(hdr_cells[5].paragraphs[0].add_run('Ograniczenia'))
            runs.append(hdr_cells[6].paragraphs[0].add_run('Unikalność'))
            runs.append(hdr_cells[7].paragraphs[0].add_run('Klucz'))
            runs.append(hdr_cells[8].paragraphs[0].add_run('Referencje'))
            runs.append(hdr_cells[9].paragraphs[0].add_run('Źródło danych'))
            for run in runs:
                run.font.size = Pt(Stage11.TABLE_FONT_SIZE)
                run.bold = True
            row_counter = 1
            for attribute in entity.attributes:
                row = table.rows[row_counter].cells
                row[0].paragraphs[0].add_run(attribute.name).font.size = Pt(Stage11.TABLE_FONT_SIZE)
                row[1].paragraphs[0].add_run(repr(attribute.type)).font.size = Pt(Stage11.TABLE_FONT_SIZE)
                if attribute.is_key:
                    row[7].paragraphs[0].add_run('PK').font.size = Pt(Stage11.TABLE_FONT_SIZE)
                row_counter += 1
            for foreign_key in entity.foreign_keys:
                row = table.rows[row_counter].cells
                row[0].paragraphs[0].add_run(foreign_key.name).font.size = Pt(Stage11.TABLE_FONT_SIZE)
                row[1].paragraphs[0].add_run(repr(foreign_key.type)).font.size = Pt(Stage11.TABLE_FONT_SIZE)
                if row[7].text == '':
                    row[7].paragraphs[0].add_run('FK').font.size = Pt(Stage11.TABLE_FONT_SIZE)
                else:
                    row[7].paragraphs[0].add_run(', FK').font.size = Pt(Stage11.TABLE_FONT_SIZE)

                reference = self.erd.get_entity_by_pk(foreign_key.name).name_plural
                if row[8].text == '':
                    row[8].paragraphs[0].add_run(reference).font.size = Pt(Stage11.TABLE_FONT_SIZE)
                else:
                    row[8].paragraphs[0].add_run(', ' + reference).font.size = Pt(Stage11.TABLE_FONT_SIZE)
                row_counter += 1

            relation_paragraph = document.add_paragraph()
            relation_paragraph.add_run().add_break()
            relation_paragraph.add_run('Zanczenie atrybutów w schemacie relacji ' + entity.name_plural)
            relation_paragraph.add_run().add_break()

            table = document.add_table(rows=len(entity.attributes)+len(entity.foreign_keys)+1, cols=2)
            hdr_row = table.rows[0].cells
            run = hdr_row[0].paragraphs[0].add_run('Nazwa atrybutu')
            run.font.size = Pt(Stage11.TABLE_FONT_SIZE)
            run.bold = True
            run = hdr_row[1].paragraphs[0].add_run('Opis')
            run.font.size = Pt(Stage11.TABLE_FONT_SIZE)
            run.bold = True

            row_counter = 1
            for attribute in entity.attributes:
                row = table.rows[row_counter].cells
                row[0].paragraphs[0].add_run(attribute.name).font.size = Pt(Stage11.TABLE_FONT_SIZE)
                row[1].paragraphs[0].add_run(attribute.description).font.size = Pt(Stage11.TABLE_FONT_SIZE)
                row_counter += 1
            for foreign_key in entity.foreign_keys:
                row = table.rows[row_counter].cells
                row[0].paragraphs[0].add_run(foreign_key.name).font.size = Pt(Stage11.TABLE_FONT_SIZE)
                row[1].paragraphs[0].add_run(foreign_key.description).font.size = Pt(Stage11.TABLE_FONT_SIZE)
                row_counter += 1

            relation_paragraph = document.add_paragraph()

            relation_paragraph.add_run().add_break()
            relation_paragraph.add_run('Przykładowe dane tabeli o schemacie relacji ' + entity.name_plural)
            relation_paragraph.add_run().add_break()

            table = document.add_table(rows=Stage11.EXAMPLES_COUNT + 1, cols=len(entity.attributes)+len(entity.foreign_keys))
            hdr_row = table.rows[0].cells

            column_counter = 0
            for attribute in entity.attributes + entity.foreign_keys:
                run = hdr_row[column_counter].paragraphs[0].add_run(attribute.name)
                run.bold = True
                run.font.size = Pt(Stage11.TABLE_FONT_SIZE)

                if attribute.is_key or attribute in entity.foreign_keys:
                    for i in range(1, Stage11.EXAMPLES_COUNT + 1):
                        row = table.rows[i].cells
                        row[column_counter].paragraphs[0].add_run(str(randrange(0, 1000))).font.size = Pt(Stage11.TABLE_FONT_SIZE)
                if attribute.type.short_name == 'Bool':
                    for i in range(1, Stage11.EXAMPLES_COUNT + 1):
                        row = table.rows[i].cells
                        row[column_counter].paragraphs[0].add_run(str(randrange(0,2))).font.size = Pt(Stage11.TABLE_FONT_SIZE)
                column_counter += 1

            relation_paragraph = document.add_paragraph()
            relation_paragraph.add_run().add_break()
            relation_paragraph.add_run().add_break()