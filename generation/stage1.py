from docx import Document
from docx.shared import Pt

from generation import Project


class Stage1:

    def __init__(self, topic='', objective='', range='', users=''):
        self.topic = topic
        self.objective = objective
        self.range = range
        self.users = users
        self.stage_number = 1

    def build(self, document):
        header = document.add_paragraph()
        header.add_run('1. Temat, cel, zakres i użytkownicy').font.size = Pt(Project.HEADER_SIZE)
        header.add_run().add_break()

        # 1.1 Temat
        secondary_header = document.add_paragraph()
        secondary_header.add_run('1.1 Temat').font.size = Pt(Project.SECONDAR_HEADER_SIZE)
        secondary_header.add_run().add_break()
        secondary_header.add_run(self.topic)
        secondary_header.add_run().add_break()

        # 1.2 Cel
        secondary_header = document.add_paragraph()
        secondary_header.add_run('1.2 Cel').font.size = Pt(Project.SECONDAR_HEADER_SIZE)
        secondary_header.add_run().add_break()
        secondary_header.add_run(self.objective)
        secondary_header.add_run().add_break()

        # 1.3 Zakres
        secondary_header = document.add_paragraph()
        secondary_header.add_run('1.3 Zakres').font.size = Pt(Project.SECONDAR_HEADER_SIZE)
        secondary_header.add_run().add_break()
        secondary_header.add_run(self.range)
        secondary_header.add_run().add_break()

        # 1.4 Użytkownicy
        secondary_header = document.add_paragraph()
        secondary_header.add_run('1.4 Użytkownicy').font.size = Pt(Project.SECONDAR_HEADER_SIZE)
        secondary_header.add_run().add_break()
        secondary_header.add_run(self.users)
        secondary_header.add_run().add_break()

        document.add_page_break()

        print('Etap 1 wygenerowany')
