from docx.shared import Pt

from generation import Project

class Stage12(object):

    def __init__(self, erd):
        self.erd = erd
        self.stage_number = 12

    def build(self, document):

        header_paragraph = document.add_paragraph()
        header_paragraph.add_run('12. Schemat bazy danych ze słownikiem atrybutów').font.size = Pt(Project.HEADER_SIZE)
        header_paragraph.add_run().add_break()

        schema_paragraph = document.add_paragraph()
        schema_paragraph.add_run('12.1 Schemat bazy danych ze słownikiem atrybutów').font.size = Pt(Project.SECONDAR_HEADER_SIZE)
        schema_paragraph.add_run().add_break()
        schema_paragraph.add_run().add_break()

        schema_paragraph.add_run('<Tytuł projektu>').bold = True
        schema_paragraph.add_run().add_break()
        schema_paragraph.add_run().add_break()

        for entity in self.erd.entities:
            schema_paragraph.add_run(entity.name_plural).bold = True
            schema_paragraph.add_run(' ')
            entity.build_argument_list(schema_paragraph)
            schema_paragraph.add_run().add_break()

        schema_paragraph.add_run().add_break()

        dictionary_paragraph = document.add_paragraph()
        dictionary_paragraph.add_run('12.2 Słownik atrybutów').font.size = Pt(Project.SECONDAR_HEADER_SIZE)
        dictionary_paragraph.add_run().add_break()

        attributes_all = []
        for entity in self.erd.entities:
            for attribute in entity.attributes:
                attributes_all.append( (attribute, entity))
            for attribute in entity.foreign_keys:
                attributes_all.append((attribute, entity))

        attributes_all.sort(key=lambda x: x[0].name)

        table = document.add_table(rows=len(attributes_all)+1, cols=3)

        hdr_row = table.rows[0].cells
        hdr_row[0].paragraphs[0].add_run('Nazwa atrybutu').bold = True
        hdr_row[1].paragraphs[0].add_run('Dziedzina atrybutu').bold = True
        hdr_row[2].paragraphs[0].add_run('Przynależność do schematu relacji').bold = True

        row_counter = 1
        for attribute in attributes_all:
            row = table.rows[row_counter].cells
            row[0].text = attribute[0].name
            row[1].text = attribute[0].type.short_name
            row[2].text = attribute[1].name_plural
            row_counter += 1
        document.add_page_break()
