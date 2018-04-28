from docx.shared import Pt


from generation import Project
from base.rule import Rule


class Stage4:

    def __init__(self, erd, rules):
        self.erd = erd
        self.rules = rules
        self.stage_number = 4

        # self.rules_dict = {}


    def build(self, document):

        header = document.add_paragraph()
        header.add_run('4. Reguły funkcjonowania').font.size = Pt(Project.HEADER_SIZE)
        header.add_run().add_break()


        for entity in self.erd.entities:
            rules_paragraph = document.add_paragraph()
            rules_paragraph.keep_together = True

            rules_paragraph.add_run(str(entity.id) + '. Reguły dla KAT/' + '{0:03}'.format(entity.id) + ' ' + entity.name_singular).font.size = Pt(16)
            rules_paragraph.add_run().add_break()
            entitys_rules = list(filter(lambda r: r.left_entity_name == entity.name_singular, self.rules))
            for rule in entitys_rules:
                rules_paragraph.add_run('REG/' + '{0:03}'.format(rule.id)).font.bold = True
                rules_paragraph.add_run('\t\t')
                rules_paragraph.add_run(repr(rule))
                rules_paragraph.add_run().add_break()

        document.add_page_break()

        print('Etap 4 wygenerowany')
