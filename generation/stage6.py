from docx.shared import Pt

from generation import Project

class Stage6(object):

    def __init__(self, erd, transactions):
        self.erd = erd
        self.transactions = transactions
        self.stage_number = 6

    def build(self, document):
        header = document.add_paragraph()
        header.add_run('6. Transakcje').font.size = Pt(Project.HEADER_SIZE)
        header.add_run().add_break()

        for transaction in self.transactions:
            transaction_paragraph = document.add_paragraph()
            transaction_paragraph.add_run('TRA/' + '{0:03}'.format(transaction.id) + '\n').bold = True
            transaction_paragraph.add_run().add_break()
            transaction_paragraph.add_run('Opis: ' + transaction.description + '\n')
            transaction_paragraph.add_run().add_break()
            transaction_paragraph.add_run('Uwarunkowania: ' + transaction.conditions + '\n')

            table = document.add_table(rows=3, cols=3)
            table.rows[0].cells[1].paragraphs[0].add_run('Wejście').bold = True
            table.rows[0].cells[2].paragraphs[0].add_run('Wyjście').bold = True
            table.rows[1].cells[0].paragraphs[0].add_run('Użytkownik').bold = True
            table.rows[2].cells[0].paragraphs[0].add_run('Baza danych').bold = True

            document.add_paragraph()

        document.add_page_break()
