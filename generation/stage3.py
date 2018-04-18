from docx.shared import Pt


from generation import Project


class Stage3:

    def __init__(self, erd):
        self.erd = erd
        self.stage_number = 3

    def build(self, document):

        header = document.add_paragraph()
        header.add_run('3. Kategorie').font.size = Pt(Project.HEADER_SIZE)
        header.add_run().add_break()

        self.erd.entities.sort(key=lambda x: x.id)

        for entity in self.erd.entities:
            entity_paragraph = document.add_paragraph()
            entity_paragraph.paragraph_format.keep_together = True
            entity_paragraph.keep_together = True
            entity_paragraph.add_run('KAT/' + '{0:03}'.format(entity.id) + ' ' + entity.name_singular).font.size = Pt(16)
            entity_paragraph.add_run().add_break()
            entity_paragraph.add_run('\tOpis: ').bold = True
            entity_paragraph.add_run('[Tutaj krótki opis encji]')
            entity_paragraph.add_run().add_break()
            entity_paragraph.add_run('\tAtrybuty:').bold = True
            entity_paragraph.add_run().add_break()
            for attribute in entity.attributes:
                entity_paragraph.add_run('\t\t' + repr(attribute.name) + ' - ' + attribute.description)
                entity_paragraph.add_run().add_break()
        document.add_page_break()

        print('Etap 3 wygenerowany')
