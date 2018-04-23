from docx.shared import Pt

from generation import Project


class Stage2:
    def __init__(self, reality_description='', dictionary='', users='', functional_requirements='', nonfunctional_requirements='',
                 existing_database='', cost=''):
        self.reality_description = reality_description
        self.dictionary = dictionary
        self.users = users
        self.functional_requirements = functional_requirements
        self.nonfunctional_requirements = nonfunctional_requirements
        self.existing_database = existing_database
        self.cost = cost
        self.stage_number = 2

    def build(self, document):
        header = document.add_paragraph()
        header.add_run('2. Analiza wycinka rzeczywistości').font.size = Pt(Project.HEADER_SIZE)

        # 2.1 Szczególowy opis wycinka rzeczywistości
        secondary_header = document.add_paragraph()
        secondary_header.add_run('1.1 Temat').font.size = Pt(Project.SECONDAR_HEADER_SIZE)
        secondary_header.add_run().add_break()
        secondary_header.add_run(self.reality_description)
        secondary_header.add_run().add_break()

        # 2.2 Słownik pojęć
        secondary_header = document.add_paragraph()
        secondary_header.add_run('2.2 Słownik pojęć').font.size = Pt(Project.SECONDAR_HEADER_SIZE)
        secondary_header.add_run().add_break()
        secondary_header.add_run(self.dictionary)
        secondary_header.add_run().add_break()

        # 2.3 Użytkownicy i zakres uprawnień
        secondary_header = document.add_paragraph()
        secondary_header.add_run('2.3 Użytkownicy i zakres uprawnień').font.size = Pt(Project.SECONDAR_HEADER_SIZE)
        secondary_header.add_run().add_break()
        secondary_header.add_run(self.users)
        secondary_header.add_run().add_break()

        # 2.4 Wymagania funkcjonalne
        secondary_header = document.add_paragraph()
        secondary_header.add_run('2.4 Wymagania funkcjonalne').font.size = Pt(Project.SECONDAR_HEADER_SIZE)
        secondary_header.add_run().add_break()
        secondary_header.add_run(self.functional_requirements)
        secondary_header.add_run().add_break()

        # 2.5 Wymagania niefunkcjonalne
        secondary_header = document.add_paragraph()
        secondary_header.add_run('2.5 Wymagania niefunkcjonalne').font.size = Pt(Project.SECONDAR_HEADER_SIZE)
        secondary_header.add_run().add_break()
        secondary_header.add_run(self.nonfunctional_requirements)
        secondary_header.add_run().add_break()

        # 2.6 Analiza istniejącej bazy danych
        secondary_header = document.add_paragraph()
        secondary_header.add_run('2.6 Analiza istniejącej bazy danych').font.size = Pt(Project.SECONDAR_HEADER_SIZE)
        secondary_header.add_run().add_break()
        secondary_header.add_run(self.existing_database)
        secondary_header.add_run().add_break()

        # 2.7 Analiza kosztów
        secondary_header = document.add_paragraph()
        secondary_header.add_run('2.7 Analiza kosztów').font.size = Pt(Project.SECONDAR_HEADER_SIZE)
        secondary_header.add_run().add_break()
        secondary_header.add_run(self.cost)
        secondary_header.add_run().add_break()


        document.add_page_break()

        print('Etap 2 wygenerowany')
