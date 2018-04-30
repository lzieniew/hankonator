from docx.shared import Pt

from generation import Project

class Stage7(object):

    def __init__(self, erd, rules):
        self.erd = erd
        self.rules = rules
        self.stage_number = 7

    def build(self, document):
        header = document.add_paragraph()
        header.add_run(u'7. Definicje encji i związków').font.size = Pt(Project.HEADER_SIZE)
        header.add_run().add_break()

        entities_paragraph = document.add_paragraph()
        entities_paragraph.add_run('7.1 Encje').font.size = Pt(Project.SECONDAR_HEADER_SIZE)
        entities_paragraph.add_run().add_break()

        for entity in self.erd.entities:
            entity_paragraph = document.add_paragraph()
            entity_paragraph.add_run('ENC/' + '{0:03}'.format(entity.id) + ' ' + entity.name_singular.upper() + '\n').bold = True
            entity_paragraph.add_run().add_break()
            entity_paragraph.add_run('\tSemantyka encji: ' + entity.description + '\n').italic = True
            entity_paragraph.add_run('\tWykaz atrybutów:\n').italic = True

            table = document.add_table(rows=len(entity.attributes) + 1, cols=4)
            table.style = 'TableGrid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].paragraphs[0].add_run('Nazwa atrybutu').bold = True
            hdr_cells[1].paragraphs[0].add_run('Opis atrybutu').bold = True
            hdr_cells[2].paragraphs[0].add_run('Typ').bold = True
            hdr_cells[3].paragraphs[0].add_run('OBL(+), OPC(-)').bold = True

            counter = 1
            for attribute in entity.attributes:
                row = table.rows[counter].cells
                row[0].text = attribute.name
                row[1].text = attribute.description
                row[2].text = repr(attribute.type)

                counter += 1

            entities_paragraph2 = document.add_paragraph()
            entities_paragraph2.add_run().add_break()
            entities_paragraph2.add_run('Klucze kandydujące:').italic = True
            entities_paragraph2.add_run(' ').add_break()
            entities_paragraph2.add_run('Klucz główny: ' + entity.get_key().name).italic = True
            entities_paragraph2.add_run(' ').add_break()
            entities_paragraph2.add_run('Charakter encji: ' + ('silna' if entity.is_strong else u'słaba')).italic = True
            entities_paragraph2.add_run(' ').add_break()
            document.add_page_break()


        relationships_paragraph = document.add_paragraph()
        relationships_paragraph.add_run('7.2 Związki').font.size = Pt(Project.SECONDAR_HEADER_SIZE)
        relationships_paragraph.add_run().add_break()

        for relationship in self.erd.relationships:
            relationship_paragraph = document.add_paragraph()
            relationship_paragraph.add_run('ZWI/' + '{0:03}'.format(relationship.id) + ' ' + relationship.name).bold\
                = True
            relationship_paragraph.add_run('(' + relationship.left_entity.upper() + '(' + relationship.left_quantity
                                           + '):' + relationship.right_entity.upper() + '('
                                           + relationship.right_quantity + ')')
            relationship_paragraph.add_run().add_break()
            entity_rules = []
            for rule in self.rules:
                if (rule.left_entity_name == relationship.left_entity \
                        or rule.left_entity_name ==  relationship.right_entity) \
                        and (rule.right_entity_name == relationship.left_entity \
                        or rule.right_entity_name == relationship.right_entity):
                    entity_rules.append(rule)
            for rule in entity_rules:
                relationship_paragraph.add_run('REG/' + '{0:03}'.format(rule.id) + ' ' + rule.content)
        print('Etap 7 wygenerowany!')
